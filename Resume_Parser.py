import os
import re
import json
import csv
import uuid
import sqlite3
import threading
import webbrowser
from collections import Counter
from datetime import datetime

import numpy as np
import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt

from flask import Flask, request, render_template_string, url_for, redirect, send_from_directory
from werkzeug.utils import secure_filename

try:
    import PyPDF2
except Exception:
    PyPDF2 = None

try:
    import docx
except Exception:
    docx = None

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except Exception:
    OCR_AVAILABLE = False


# ===================== 全局配置 =====================
plt.rcParams["font.sans-serif"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

UPLOAD_FOLDER = "uploads"
STATIC_FOLDER = "static"
EXPORT_FOLDER = "exports"
DB_PATH = "resume_parser.db"
ALLOWED_EXTENSIONS = {"pdf", "txt", "docx"}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(STATIC_FOLDER, exist_ok=True)
os.makedirs(EXPORT_FOLDER, exist_ok=True)


# ===================== 工具函数 =====================
def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[-1].lower() in ALLOWED_EXTENSIONS


def unique_list(items):
    seen = set()
    result = []
    for item in items:
        item = str(item or "").strip()
        if not item:
            continue
        key = item.lower()
        if key not in seen:
            seen.add(key)
            result.append(item)
    return result


def clean_line(line: str) -> str:
    line = str(line or "")
    line = line.replace("\xa0", " ").replace("\u3000", " ")
    line = re.sub(r"[\t\r\f\v]+", " ", line)
    line = re.sub(r"\s{2,}", " ", line)
    return line.strip(" |｜\n")


def normalize_token(text: str) -> str:
    return str(text or "").strip().lower().replace(" ", "").replace("．", ".")


def safe_original_filename(filename: str) -> str:
    return os.path.basename(filename or "").strip()


def text_density_score(line: str) -> int:
    bad_keywords = [
        "简历", "个人", "求职", "应聘", "岗位", "电话", "手机", "邮箱", "地址", "教育", "经历", "技能",
        "项目", "实习", "工作", "校园", "证书", "荣誉", "自我", "评价", "专业", "学校", "学院", "大学",
        "民族", "政治", "出生", "年龄", "性别", "籍贯", "现居", "期望", "薪资", "照片", "基本信息",
        "resume", "cv", "profile", "objective", "education", "experience", "skill", "email", "phone",
    ]
    score = 0
    lower = line.lower()
    for kw in bad_keywords:
        if kw in lower:
            score += 6
    if re.search(r"\d", line):
        score += 4
    if re.search(r"[@:/\\]", line):
        score += 5
    if len(line) > 12:
        score += 3
    return score


# ===================== 数据类 =====================
class Resume:
    def __init__(self, filename, original_filename=None):
        self.filename = filename
        self.original_filename = original_filename or filename
        self.name = ""
        self.phone = ""
        self.email = ""
        self.education = []
        self.skills = []
        self.raw_preview = ""
        self.source_stats = {
            "chars": 0,
            "tables": 0,
            "ocr_used": False,
            "parser": "",
            "warning": "",
        }

    def to_dict(self):
        return {
            "filename": self.filename,
            "original_filename": self.original_filename,
            "name": self.name,
            "phone": self.phone,
            "email": self.email,
            "education": self.education,
            "skills": self.skills,
            "raw_preview": self.raw_preview,
            "source_stats": self.source_stats,
        }


# ===================== 简历解析类 =====================
class Parser:
    PHONE_PATTERN = re.compile(r"(?:\+?86[-\s]?)?1[3-9]\d{9}")
    EMAIL_PATTERN = re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}")

    EDU_KEYWORDS = [
        "博士", "硕士", "研究生", "本科", "学士", "大专", "专科", "高中",
        "大学", "学院", "学校", "专业", "教育经历", "教育背景", "学历",
    ]

    STOP_SECTION_HEADERS = [
        "工作经历", "项目经历", "实习经历", "校园经历", "获奖", "证书", "自我评价", "个人评价",
        "求职意向", "联系方式", "个人信息", "兴趣爱好", "语言能力", "荣誉", "培训经历",
    ]

    SKILL_LIBRARY = {
        # 计算机 / 软件
        "Python": ["Python", "py"],
        "Java": ["Java"],
        "C": ["C语言"],
        "C++": ["C++", "CPP"],
        "C#": ["C#"],
        "Go": ["Golang", "Go语言"],
        "SQL": ["SQL"],
        "MySQL": ["MySQL"],
        "SQLite": ["SQLite"],
        "Redis": ["Redis"],
        "MongoDB": ["MongoDB"],
        "Flask": ["Flask"],
        "Django": ["Django"],
        "FastAPI": ["FastAPI"],
        "Spring": ["Spring"],
        "SpringBoot": ["SpringBoot", "Spring Boot"],
        "Maven": ["Maven"],
        "HTML": ["HTML", "HTML5"],
        "CSS": ["CSS", "CSS3"],
        "JavaScript": ["JavaScript", "JS"],
        "TypeScript": ["TypeScript", "TS"],
        "Vue": ["Vue", "Vue.js", "Vue2", "Vue3"],
        "React": ["React", "React.js"],
        "Node.js": ["Node.js", "NodeJS", "Node"],
        "Linux": ["Linux"],
        "Shell": ["Shell", "Bash"],
        "Git": ["Git", "GitHub", "GitLab"],
        "Docker": ["Docker"],
        "Kubernetes": ["Kubernetes", "K8s"],
        "Nginx": ["Nginx"],
        "Jenkins": ["Jenkins"],
        "Selenium": ["Selenium"],
        "Postman": ["Postman"],
        "JMeter": ["JMeter", "Jmeter"],
        "网络安全": ["网络安全", "信息安全"],
        "渗透测试": ["渗透测试"],
        "防火墙": ["防火墙"],
        "SQL注入": ["SQL注入"],
        "漏洞挖掘": ["漏洞挖掘"],

        # 数据 / 算法
        "NumPy": ["NumPy", "numpy"],
        "Pandas": ["Pandas", "pandas"],
        "Matplotlib": ["Matplotlib", "matplotlib"],
        "Excel": ["Excel", "WPS表格"],
        "Tableau": ["Tableau"],
        "PowerBI": ["PowerBI", "Power BI"],
        "Hadoop": ["Hadoop"],
        "Spark": ["Spark"],
        "Hive": ["Hive"],
        "Kafka": ["Kafka"],
        "Scala": ["Scala"],
        "数据结构": ["数据结构"],
        "算法": ["算法"],
        "多线程": ["多线程"],
        "机器学习": ["机器学习", "Machine Learning", "ML"],
        "深度学习": ["深度学习", "Deep Learning", "DL"],
        "爬虫": ["爬虫", "Scrapy", "BeautifulSoup", "Requests"],
        "数据分析": ["数据分析"],
        "数据可视化": ["数据可视化", "可视化"],

        # 机械 / 电气 / 自动化 / 通信 / 电子
        "机械设计": ["机械设计", "结构设计", "机械制图"],
        "机械原理": ["机械原理"],
        "机械制造": ["机械制造", "制造工艺", "机加工"],
        "CAD": ["CAD", "AutoCAD"],
        "SolidWorks": ["SolidWorks", "SW"],
        "UG": ["UG", "NX"],
        "CATIA": ["CATIA"],
        "Creo": ["Creo", "ProE", "Pro/E"],
        "PLC": ["PLC", "西门子PLC", "三菱PLC"],
        "电气控制": ["电气控制", "电控"],
        "电路设计": ["电路设计", "模拟电路", "数字电路"],
        "单片机": ["单片机", "MCU"],
        "嵌入式": ["嵌入式", "嵌入式开发"],
        "ARM": ["ARM"],
        "STM32": ["STM32"],
        "PCB": ["PCB", "Altium Designer", "AD软件"],
        "传感器": ["传感器"],
        "自动化": ["自动化", "自动控制"],
        "控制理论": ["控制理论", "PID", "现代控制", "智能控制"],
        "MATLAB": ["MATLAB", "Simulink"],
        "LabVIEW": ["LabVIEW"],
        "ROS": ["ROS", "ROS2"],
        "机器人": ["机器人", "机械臂", "移动机器人"],
        "通信原理": ["通信原理"],
        "信号处理": ["信号处理", "数字信号处理", "DSP"],
        "射频": ["射频", "RF"],
        "FPGA": ["FPGA", "Verilog", "VHDL"],
        "RTOS": ["RTOS"],

        # 产品 / 运营 / 设计
        "需求分析": ["需求分析"],
        "项目管理": ["项目管理"],
        "原型设计": ["原型设计"],
        "Axure": ["Axure"],
        "Figma": ["Figma"],
        "墨刀": ["墨刀"],
        "Photoshop": ["Photoshop", "PS"],
        "Illustrator": ["Illustrator", "AI软件"],
        "Sketch": ["Sketch"],
        "动效设计": ["动效设计"],
        "用户研究": ["用户研究", "用户调研"],
        "竞品分析": ["竞品分析"],
        "活动策划": ["活动策划"],
        "内容运营": ["内容运营"],
        "用户运营": ["用户运营"],
        "新媒体运营": ["新媒体运营", "新媒体", "公众号", "小红书", "抖音运营"],
        "数据运营": ["数据运营"],
        "SEO": ["SEO"],
        "SEM": ["SEM"],

        # 财务 / 人力 / 行政 / 市场 / 销售 / 物流
        "财务分析": ["财务分析"],
        "会计": ["会计", "财务会计"],
        "出纳": ["出纳"],
        "税务": ["税务", "纳税申报"],
        "审计": ["审计"],
        "成本核算": ["成本核算"],
        "招聘": ["招聘"],
        "绩效": ["绩效", "绩效考核"],
        "薪酬": ["薪酬", "薪资"],
        "员工关系": ["员工关系"],
        "培训": ["培训"],
        "行政管理": ["行政管理", "办公室管理"],
        "公文写作": ["公文写作"],
        "档案管理": ["档案管理"],
        "合同管理": ["合同管理"],
        "客户沟通": ["客户沟通", "客户维护"],
        "销售": ["销售", "销售技巧"],
        "渠道开发": ["渠道开发", "渠道拓展"],
        "商务谈判": ["商务谈判", "谈判"],
        "市场调研": ["市场调研"],
        "品牌推广": ["品牌推广", "品牌策划"],
        "供应链": ["供应链"],
        "采购": ["采购"],
        "物流": ["物流"],
        "仓储": ["仓储"],
        "质量管理": ["质量管理", "品质管理"],
        "ISO9001": ["ISO9001", "质量体系"],
        "六西格玛": ["六西格玛", "Six Sigma"],
        "生产管理": ["生产管理"],
        "精益生产": ["精益生产", "Lean"],

        # 教育 / 语言 / 通用能力
        "英语": ["英语", "CET4", "CET6", "四级", "六级"],
        "日语": ["日语", "N1", "N2"],
        "汉语言": ["汉语言", "汉语言文学", "中文"],
        "新闻": ["新闻", "新闻学", "新闻传播"],
        "教师资格证": ["教师资格证", "教资"],
        "课程设计": ["课程设计", "教学设计"],
        "授课": ["授课", "讲课", "教学"],
        "班级管理": ["班级管理"],
        "Office": ["Office", "Word", "PPT", "PowerPoint"],
        "沟通能力": ["沟通能力", "沟通协调"],
        "团队协作": ["团队协作", "团队合作"],
    }

    COMMON_CHINESE_SURNAMES = set(
        "赵钱孙李周吴郑王冯陈褚卫蒋沈韩杨朱秦尤许何吕施张孔曹严华金魏陶姜"
        "戚谢邹喻柏水窦章云苏潘葛奚范彭郎鲁韦昌马苗凤花方俞任袁柳鲍史唐"
        "费廉岑薛雷贺倪汤滕殷罗毕郝邬安常乐于时傅皮卞齐康伍余元卜顾孟"
        "平黄和穆萧尹姚邵湛汪祁毛禹狄米贝明臧计伏成戴谈宋庞熊纪舒屈项"
        "祝董梁杜阮蓝闵席季麻强贾路娄危江童颜郭梅盛林刁钟徐邱骆高夏蔡田"
        "胡凌霍虞万支柯昝管卢莫经房裘缪干解应宗丁宣邓郁单杭洪包诸左石崔"
        "吉龚程邢裴陆荣翁荀羊於惠甄曲家封芮羿储靳汲邴糜松井段富巫乌焦巴"
        "弓牧隗山谷车侯宓蓬全郗班仰秋仲伊宫宁仇栾暴甘斜厉戎祖武符刘景詹"
        "束龙叶幸司韶郜黎蓟薄印宿白怀蒲台从鄂索咸籍赖卓蔺屠蒙池乔阴郁胥"
        "能苍双闻莘党翟谭贡劳逄姬申扶堵冉宰郦雍却璩桑桂濮牛寿通边扈燕冀"
        "浦尚农温别庄晏柴瞿阎充慕连茹习宦艾鱼容向古易慎戈廖庾终暨居衡步"
        "都耿满弘匡国文寇广禄阙东欧殳沃利蔚越夔隆师巩厍聂晁勾敖融冷訾辛"
        "阚那简饶空曾毋沙乜养鞠须丰巢关蒯相查后荆红游竺权逯盖益桓公"
    )

    COMPOUND_SURNAMES = [
        "欧阳", "太史", "端木", "上官", "司马", "东方", "独孤", "南宫", "万俟", "闻人", "夏侯",
        "诸葛", "尉迟", "公羊", "赫连", "澹台", "皇甫", "宗政", "濮阳", "公冶", "太叔", "申屠",
        "公孙", "慕容", "仲孙", "钟离", "长孙", "宇文", "司徒", "鲜于", "司空", "闾丘", "子车",
        "亓官", "司寇", "巫马", "公西", "颛孙", "壤驷", "公良", "漆雕", "乐正", "宰父", "谷梁",
    ]

    def __init__(self, resume_dir, filename_map=None):
        self.resume_dir = resume_dir
        self.filename_map = filename_map or {}
        self.resumes = []

    def read_file(self, filepath):
        ext = filepath.rsplit(".", 1)[-1].lower()
        if ext == "pdf":
            return self._read_pdf(filepath)
        if ext == "docx":
            return self._read_docx(filepath)
        if ext == "txt":
            return self._read_txt(filepath)
        return "", {"tables": 0, "ocr_used": False, "parser": "unsupported", "warning": "不支持的文件类型"}

    def _read_txt(self, filepath):
        for enc in ["utf-8", "utf-8-sig", "gbk", "gb18030"]:
            try:
                with open(filepath, "r", encoding=enc, errors="ignore") as f:
                    return self._clean_text(f.read()), {
                        "tables": 0,
                        "ocr_used": False,
                        "parser": f"txt:{enc}",
                        "warning": "",
                    }
            except Exception:
                continue
        return "", {"tables": 0, "ocr_used": False, "parser": "txt", "warning": "TXT读取失败"}

    def _read_pdf(self, filepath):
        text_parts = []
        table_count = 0
        parser_names = []
        warning = ""

        if pdfplumber is not None:
            try:
                parser_names.append("pdfplumber")
                with pdfplumber.open(filepath) as pdf:
                    for page_no, page in enumerate(pdf.pages, start=1):
                        page_text = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
                        if page_text.strip():
                            text_parts.append(f"\n【PDF第{page_no}页正文】\n{page_text}")

                        tables = page.extract_tables() or []
                        for t_idx, table in enumerate(tables, start=1):
                            table_text = self._table_to_text(table)
                            if table_text:
                                table_count += 1
                                text_parts.append(f"\n【PDF第{page_no}页表格{t_idx}】\n{table_text}")
            except Exception as e:
                warning += f"pdfplumber失败：{e}; "
        else:
            warning += "未安装pdfplumber，PDF表格提取能力受限; "

        if PyPDF2 is not None:
            try:
                parser_names.append("PyPDF2")
                with open(filepath, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page_no, page in enumerate(reader.pages, start=1):
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            text_parts.append(f"\n【PyPDF2第{page_no}页】\n{page_text}")
            except Exception as e:
                warning += f"PyPDF2失败：{e}; "
        else:
            warning += "未安装PyPDF2; "

        text = self._clean_text("\n".join(text_parts))
        ocr_used = False

        if len(text) < 80:
            if OCR_AVAILABLE:
                try:
                    parser_names.append("OCR")
                    images = convert_from_path(filepath, dpi=220)
                    ocr_parts = []
                    for i, img in enumerate(images, start=1):
                        ocr_text = pytesseract.image_to_string(img, lang="chi_sim+eng")
                        if ocr_text.strip():
                            ocr_parts.append(f"\n【OCR第{i}页】\n{ocr_text}")
                    ocr_text_all = self._clean_text("\n".join(ocr_parts))
                    if ocr_text_all:
                        text = ocr_text_all
                        ocr_used = True
                    else:
                        warning += "OCR未识别到有效文本; "
                except Exception as e:
                    warning += f"OCR失败：{e}; "
            else:
                warning += "文本很少，可能是扫描版PDF；如需识别请安装OCR组件; "

        return text, {
            "tables": table_count,
            "ocr_used": ocr_used,
            "parser": "+".join(unique_list(parser_names)) or "pdf",
            "warning": warning.strip(),
        }

    def _read_docx(self, filepath):
        text_parts = []
        table_count = 0
        warning = ""

        if docx is None:
            return "", {"tables": 0, "ocr_used": False, "parser": "python-docx", "warning": "未安装python-docx"}

        try:
            doc = docx.Document(filepath)

            for para in doc.paragraphs:
                txt = clean_line(para.text)
                if txt:
                    text_parts.append(txt)

            for t_idx, table in enumerate(doc.tables, start=1):
                table_text = self._docx_table_to_text(table)
                if table_text:
                    table_count += 1
                    text_parts.append(f"\n【Word表格{t_idx}】\n{table_text}")

            for section in doc.sections:
                for para in section.header.paragraphs:
                    txt = clean_line(para.text)
                    if txt:
                        text_parts.append(f"【页眉】{txt}")
                for para in section.footer.paragraphs:
                    txt = clean_line(para.text)
                    if txt:
                        text_parts.append(f"【页脚】{txt}")
        except Exception as e:
            warning = f"DOCX读取失败：{e}"

        return self._clean_text("\n".join(text_parts)), {
            "tables": table_count,
            "ocr_used": False,
            "parser": "python-docx",
            "warning": warning,
        }

    def _table_to_text(self, table):
        rows = []
        for row in table:
            if not row:
                continue
            cells = [clean_line(cell) for cell in row]
            cells = [c for c in cells if c]
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _docx_table_to_text(self, table):
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = clean_line(" ".join(p.text for p in cell.paragraphs if p.text))
                if cell_text:
                    cells.append(cell_text)
            cells = unique_list(cells)
            if cells:
                rows.append(" | ".join(cells))
        return "\n".join(rows)

    def _clean_text(self, text):
        text = str(text or "").replace("\xa0", " ").replace("\u3000", " ")
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    def extract_info(self, content, original_filename=""):
        data = {"phone": "", "email": "", "education": [], "skills": [], "name": ""}

        phone_match = self.PHONE_PATTERN.search(content)
        if phone_match:
            data["phone"] = re.sub(r"\D", "", phone_match.group())[-11:]

        email_match = self.EMAIL_PATTERN.search(content)
        if email_match:
            data["email"] = email_match.group()

        data["name"] = self._extract_name(content, original_filename=original_filename)
        data["education"] = self._extract_education(content)
        data["skills"] = self._extract_skills(content)

        return data

    def _extract_name(self, content, original_filename=""):
        lines = [clean_line(x) for x in content.split("\n") if clean_line(x)]

        explicit_patterns = [
            r"(?:姓名|姓\s*名|Name|name)\s*[：:|｜]?\s*([\u4e00-\u9fa5]{2,4})",
            r"(?:姓名|姓\s*名|Name|name)\s*[：:|｜]?\s*([A-Za-z][A-Za-z\s]{1,30})",
        ]

        for p in explicit_patterns:
            m = re.search(p, content, flags=re.I)
            if m:
                candidate = clean_line(m.group(1))
                if self._is_probable_name(candidate):
                    return candidate

        table_like_patterns = [
            r"(?:^|\n|[|｜])\s*([\u4e00-\u9fa5]{2,4})\s*(?:[|｜,，\s]+)(?:男|女)(?:[|｜,，\s]+|$)",
            r"(?:^|\n)\s*([\u4e00-\u9fa5]{2,4})\s+(?:男|女)\s+(?:\d{2}\s*岁?|党员|团员|群众|本科|硕士|大专)",
            r"(?:^|\n)\s*([\u4e00-\u9fa5]{2,4})\s+(?:\+?86[-\s]?)?1[3-9]\d{9}",
        ]

        for p in table_like_patterns:
            m = re.search(p, content)
            if m:
                candidate = clean_line(m.group(1))
                if self._is_probable_name(candidate):
                    return candidate

        top_lines = self._get_top_content_lines(lines, limit=28)
        scored = []

        for idx, line in enumerate(top_lines):
            candidate = self._extract_name_from_line(line)
            if candidate and self._is_probable_name(candidate):
                score = 100
                score -= idx * 3
                score -= text_density_score(line)

                if line == candidate:
                    score += 35

                nearby = " ".join(top_lines[max(0, idx - 3): idx + 4])
                if self.PHONE_PATTERN.search(nearby) or self.EMAIL_PATTERN.search(nearby):
                    score += 18
                if re.search(r"男|女|年龄|出生|籍贯|民族|政治面貌", nearby):
                    score += 10

                scored.append((score, candidate))

        if scored:
            scored.sort(reverse=True, key=lambda x: x[0])
            if scored[0][0] >= 65:
                return scored[0][1]

        fname = os.path.splitext(os.path.basename(original_filename or ""))[0]
        fname = re.sub(
            r"(?i)resume|cv|简历|个人|应聘|求职|前端|后端|开发|工程师|产品|测试|数据|分析|设计|岗位|\d+",
            " ",
            fname,
        )

        for token in re.split(r"[-_—\s（）()【】\[\]]+", fname):
            token = clean_line(token)
            if self._is_probable_name(token):
                return token

        return ""

    def _get_top_content_lines(self, lines, limit=28):
        result = []
        for line in lines:
            if re.match(r"^【.*?(PDF第|PyPDF2第|Word表格|OCR第|页眉|页脚).*?】$", line):
                continue
            if line in {"个人简历", "简历", "RESUME", "Resume", "个人信息", "基本信息"}:
                continue
            result.append(line)
            if len(result) >= limit:
                break
        return result

    def _extract_name_from_line(self, line):
        line = clean_line(line)
        if not line:
            return ""

        cleaned = re.sub(r"^(应聘者|候选人|个人信息|基本信息|姓名)\s*[：:]?\s*", "", line)
        cleaned = clean_line(cleaned)

        if re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", cleaned):
            return cleaned

        no_space = re.sub(r"\s+", "", cleaned)
        if re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", no_space):
            return no_space

        m = re.match(
            r"^([\u4e00-\u9fa5]{2,4})(?:\s+|[|｜,，])(?:男|女|\d{2}\s*岁?|本科|硕士|大专|党员|团员|群众|1[3-9]\d{9})",
            cleaned,
        )
        if m:
            return m.group(1)

        if len(cleaned) <= 30 and re.fullmatch(r"[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,2}", cleaned):
            if cleaned.lower() not in {"java", "python", "react", "vue", "resume", "profile"}:
                return cleaned

        return ""

    def _is_probable_name(self, candidate):
        candidate = clean_line(candidate)
        if not candidate:
            return False

        reject_words = [
            "个人", "简历", "求职", "应聘", "岗位", "电话", "手机", "邮箱", "地址", "教育", "经历", "技能", "项目",
            "学校", "学院", "大学", "专业", "本科", "硕士", "博士", "大专", "男", "女", "党员", "团员", "群众",
            "北京", "上海", "广州", "深圳", "杭州", "南京", "武汉", "成都", "重庆", "西安", "天津", "苏州",
        ]

        if any(w in candidate for w in reject_words):
            return False
        if self.PHONE_PATTERN.search(candidate) or self.EMAIL_PATTERN.search(candidate):
            return False
        if re.search(r"\d|[@:/\\]|经验|开发|工程师|产品|运营|测试|设计", candidate):
            return False

        if re.fullmatch(r"[\u4e00-\u9fa5]{2,4}", candidate):
            if any(candidate.startswith(s) and len(candidate) > len(s) for s in self.COMPOUND_SURNAMES):
                return True
            return candidate[0] in self.COMMON_CHINESE_SURNAMES

        if re.fullmatch(r"[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+){0,2}", candidate):
            return True

        return False

    def _extract_education(self, content):
        lines = [clean_line(x) for x in content.split("\n") if clean_line(x)]
        candidates = []

        edu_block = self._extract_section(content, ["教育经历", "教育背景", "学历教育", "教育"])
        edu_lines = [clean_line(x) for x in edu_block.split("\n") if clean_line(x)] if edu_block else []

        year_pattern = re.compile(r"(19\d{2}|20\d{2}).{0,20}(19\d{2}|20\d{2}|至今|现在)")

        for line in edu_lines + lines:
            if len(line) > 180:
                continue

            has_edu = any(k in line for k in self.EDU_KEYWORDS)
            has_year = bool(year_pattern.search(line))

            if has_edu or (has_year and any(x in line for x in ["专业", "大学", "学院", "本科", "硕士", "大专"])):
                if line not in ["教育", "教育经历", "教育背景", "学历"]:
                    candidates.append(line)

        return unique_list(candidates)[:10]

    def _extract_section(self, content, headers):
        lines = content.split("\n")
        start = None

        for i, line in enumerate(lines):
            c = clean_line(line)
            if any(h in c for h in headers):
                start = i
                break

        if start is None:
            return ""

        collected = []

        for line in lines[start:start + 45]:
            c = clean_line(line)
            if collected and any(h in c for h in self.STOP_SECTION_HEADERS):
                break
            collected.append(line)

        return "\n".join(collected)

    def _extract_skills(self, content):
        candidates = []
        skill_section = self._extract_section(content, ["专业技能", "技能特长", "个人技能", "技能", "掌握技能", "IT技能"])
        scan_text = content + "\n" + skill_section

        for canonical, aliases in self.SKILL_LIBRARY.items():
            for alias in aliases + [canonical]:
                if self._contains_skill(scan_text, alias):
                    candidates.append(canonical)
                    break

        patterns = [
            r"(?:专业技能|技能特长|个人技能|技能|掌握|熟悉|精通|擅长|Skills?)\s*[：:]\s*([^\n。；;]{2,260})",
            r"(?:技术栈|开发环境|工具|软件|证书|能力)\s*[：:]\s*([^\n。；;]{2,260})",
        ]

        for p in patterns:
            for m in re.finditer(p, content, flags=re.I):
                raw = m.group(1)
                for part in re.split(r"[、，,;/；|｜\s]+", raw):
                    mapped = self._map_to_canonical_skill(part)
                    if mapped:
                        candidates.append(mapped)

        if skill_section:
            for line in skill_section.split("\n"):
                for part in re.split(r"[、，,;/；|｜]+", line):
                    mapped = self._map_to_canonical_skill(part)
                    if mapped:
                        candidates.append(mapped)

        return unique_list(candidates)

    def _contains_skill(self, text, alias):
        alias = str(alias or "").strip()
        if not alias:
            return False

        if re.search(r"[A-Za-z+#.]", alias):
            pattern = r"(?<![A-Za-z0-9_+#.])" + re.escape(alias) + r"(?![A-Za-z0-9_+#.])"
            return re.search(pattern, text, flags=re.I) is not None

        return alias in text

    def _map_to_canonical_skill(self, token):
        token = clean_line(token)
        if not token or len(token) > 35:
            return ""

        token_norm = normalize_token(token)

        for canonical, aliases in self.SKILL_LIBRARY.items():
            if token_norm == normalize_token(canonical):
                return canonical
            for alias in aliases:
                if token_norm == normalize_token(alias):
                    return canonical

        return ""

    def batch_parse(self):
        self.resumes = []
        print(f"\n开始解析文件夹: {self.resume_dir}")

        for filename in os.listdir(self.resume_dir):
            filepath = os.path.join(self.resume_dir, filename)

            if not os.path.isfile(filepath):
                continue

            ext = filename.rsplit(".", 1)[-1].lower()

            if ext not in ALLOWED_EXTENSIONS:
                print(f"跳过不支持文件: {filename}")
                continue

            original_filename = self.filename_map.get(filename, filename)
            resume = Resume(filename=filename, original_filename=original_filename)

            try:
                content, meta = self.read_file(filepath)

                resume.source_stats = {
                    "chars": len(content or ""),
                    "tables": meta.get("tables", 0),
                    "ocr_used": meta.get("ocr_used", False),
                    "parser": meta.get("parser", ""),
                    "warning": meta.get("warning", ""),
                }

                resume.raw_preview = (content or "")[:500]

                if not content:
                    resume.source_stats["warning"] = (
                        resume.source_stats.get("warning")
                        or "未提取到文本，可能是扫描版PDF、加密文档或图片型文档"
                    )
                    self.resumes.append(resume)
                    print(f"未提取到内容: {original_filename}")
                    continue

                info = self.extract_info(content, original_filename=original_filename)

                resume.name = info["name"]
                resume.phone = info["phone"]
                resume.email = info["email"]
                resume.education = info["education"]
                resume.skills = info["skills"]

                self.resumes.append(resume)

                print(
                    f"解析完成: {original_filename} | "
                    f"姓名: {resume.name or '未识别'} | "
                    f"技能数: {len(resume.skills)} | "
                    f"表格: {resume.source_stats['tables']}"
                )

            except Exception as e:
                resume.source_stats["warning"] = f"解析异常：{e}"
                self.resumes.append(resume)
                print(f"解析异常 {original_filename}: {e}")

        print(f"\n解析完成，共处理 {len(self.resumes)} 份简历")
        return self.resumes


# ===================== 报告生成类 =====================
class ReportGenerator:
    """
    v3修改重点：
    1. 岗位推荐表不显示 0 分岗位。
    2. 若某候选人所有岗位都是 0 分，则该候选人在推荐表中不展示，避免出现误导性的 0 分可培养。
    3. 可视化图表改成上下两行展示，不再左右并排。
    4. 热力图仍可包含 0 分用于整体分析，但推荐表过滤 0 分。
    """

    MIN_DISPLAY_SCORE = 1.0

    JOB_PROFILES = {
        "Python后端开发": {
            "category": "计算机/软件",
            "core": ["Python", "SQL"],
            "related": ["Flask", "Django", "FastAPI", "MySQL", "Redis", "Linux", "Git"],
            "bonus": ["Docker", "Nginx", "数据结构", "算法", "多线程"],
            "education_keywords": ["计算机", "软件", "网络工程", "信息管理", "人工智能"],
        },
        "Java后端开发": {
            "category": "计算机/软件",
            "core": ["Java", "MySQL"],
            "related": ["Spring", "SpringBoot", "Redis", "Maven", "Git", "Linux"],
            "bonus": ["Docker", "Nginx", "数据结构", "算法", "多线程"],
            "education_keywords": ["计算机", "软件", "网络工程", "信息管理"],
        },
        "前端开发": {
            "category": "计算机/软件",
            "core": ["HTML", "CSS", "JavaScript"],
            "related": ["Vue", "React", "Node.js", "TypeScript", "Git"],
            "bonus": ["Figma", "原型设计", "数据可视化"],
            "education_keywords": ["计算机", "软件", "数字媒体", "网络工程"],
        },
        "测试工程师": {
            "category": "计算机/软件",
            "core": ["Postman", "SQL"],
            "related": ["Python", "Selenium", "JMeter", "Jenkins", "Linux", "Git"],
            "bonus": ["Docker", "Java", "自动化"],
            "education_keywords": ["计算机", "软件", "网络工程", "信息管理"],
        },
        "数据分析师": {
            "category": "数据/运营",
            "core": ["Excel", "SQL", "数据分析"],
            "related": ["Python", "Pandas", "NumPy", "Matplotlib", "Tableau", "PowerBI"],
            "bonus": ["机器学习", "数据可视化", "爬虫"],
            "education_keywords": ["统计", "数学", "经济", "金融", "信息管理", "计算机", "数据科学"],
        },
        "大数据开发": {
            "category": "计算机/数据",
            "core": ["SQL", "Hadoop"],
            "related": ["Spark", "Hive", "Kafka", "Python", "Java", "Scala", "Linux"],
            "bonus": ["Docker", "机器学习"],
            "education_keywords": ["计算机", "软件", "大数据", "数据科学"],
        },
        "网络安全工程师": {
            "category": "计算机/安全",
            "core": ["Linux", "网络安全"],
            "related": ["Python", "渗透测试", "防火墙", "SQL注入", "漏洞挖掘"],
            "bonus": ["Shell", "Docker", "MySQL"],
            "education_keywords": ["网络安全", "信息安全", "计算机", "网络工程"],
        },
        "机械设计工程师": {
            "category": "机械/制造",
            "core": ["机械设计", "CAD"],
            "related": ["SolidWorks", "UG", "CATIA", "Creo", "机械原理", "机械制造"],
            "bonus": ["项目管理", "质量管理", "Office"],
            "education_keywords": ["机械", "机电", "车辆工程", "机械设计制造", "材料成型"],
        },
        "工艺工程师": {
            "category": "机械/制造",
            "core": ["机械制造", "生产管理"],
            "related": ["CAD", "机械设计", "质量管理", "精益生产", "ISO9001"],
            "bonus": ["六西格玛", "项目管理", "Office"],
            "education_keywords": ["机械", "工业工程", "材料", "制造", "机电"],
        },
        "质量工程师": {
            "category": "质量/制造",
            "core": ["质量管理", "ISO9001"],
            "related": ["六西格玛", "生产管理", "Excel", "机械制造", "精益生产"],
            "bonus": ["项目管理", "沟通能力", "Office"],
            "education_keywords": ["质量管理", "机械", "工业工程", "材料", "制造"],
        },
        "电气工程师": {
            "category": "电气/自动化",
            "core": ["电气控制", "PLC"],
            "related": ["电路设计", "自动化", "传感器", "CAD", "MATLAB"],
            "bonus": ["LabVIEW", "项目管理", "Office"],
            "education_keywords": ["电气", "自动化", "电气工程", "测控", "机电"],
        },
        "自动化工程师": {
            "category": "电气/自动化",
            "core": ["自动化", "PLC"],
            "related": ["电气控制", "控制理论", "MATLAB", "传感器", "机器人"],
            "bonus": ["Python", "ROS", "LabVIEW"],
            "education_keywords": ["自动化", "控制", "测控", "电气", "机器人工程"],
        },
        "嵌入式工程师": {
            "category": "电子/嵌入式",
            "core": ["C", "单片机"],
            "related": ["嵌入式", "STM32", "ARM", "PCB", "电路设计", "Linux"],
            "bonus": ["C++", "RTOS", "信号处理"],
            "education_keywords": ["电子", "通信", "自动化", "测控", "计算机", "物联网"],
        },
        "电子硬件工程师": {
            "category": "电子/硬件",
            "core": ["电路设计", "PCB"],
            "related": ["单片机", "STM32", "传感器", "FPGA", "信号处理"],
            "bonus": ["嵌入式", "MATLAB", "质量管理"],
            "education_keywords": ["电子", "通信", "微电子", "测控", "自动化"],
        },
        "通信工程师": {
            "category": "通信/电子",
            "core": ["通信原理", "信号处理"],
            "related": ["MATLAB", "射频", "FPGA", "Python", "Linux"],
            "bonus": ["机器学习", "数据分析", "项目管理"],
            "education_keywords": ["通信", "电子信息", "信息工程", "信号处理"],
        },
        "机器人工程师": {
            "category": "机器人/控制",
            "core": ["机器人", "控制理论"],
            "related": ["MATLAB", "ROS", "Python", "C++", "自动化", "传感器"],
            "bonus": ["机器学习", "Linux", "机械设计"],
            "education_keywords": ["机器人", "自动化", "控制", "机械电子", "测控", "计算机"],
        },
        "产品经理": {
            "category": "产品/运营",
            "core": ["需求分析", "原型设计"],
            "related": ["Axure", "Figma", "墨刀", "用户研究", "竞品分析", "项目管理", "数据分析"],
            "bonus": ["SQL", "PowerBI", "Tableau"],
            "education_keywords": ["信息管理", "计算机", "电子商务", "工商管理", "工业设计"],
        },
        "新媒体运营": {
            "category": "运营/市场",
            "core": ["新媒体运营", "内容运营"],
            "related": ["活动策划", "用户运营", "数据运营", "Excel", "Photoshop"],
            "bonus": ["SEO", "SEM", "品牌推广", "数据分析"],
            "education_keywords": ["新闻", "传播", "广告", "市场营销", "中文", "汉语言", "电子商务"],
        },
        "电商运营": {
            "category": "运营/市场",
            "core": ["数据运营", "活动策划"],
            "related": ["Excel", "用户运营", "内容运营", "客户沟通", "数据分析"],
            "bonus": ["SEO", "SEM", "Photoshop", "PowerBI"],
            "education_keywords": ["电子商务", "市场营销", "工商管理", "物流管理"],
        },
        "UI设计师": {
            "category": "设计",
            "core": ["Photoshop", "Figma"],
            "related": ["Illustrator", "Sketch", "原型设计", "动效设计"],
            "bonus": ["HTML", "CSS", "Axure"],
            "education_keywords": ["设计", "视觉传达", "数字媒体", "美术", "工业设计"],
        },
        "财务专员": {
            "category": "财务/会计",
            "core": ["会计", "Excel"],
            "related": ["财务分析", "出纳", "税务", "成本核算", "审计"],
            "bonus": ["Office", "数据分析", "沟通能力"],
            "education_keywords": ["会计", "财务", "审计", "金融", "经济"],
        },
        "人力资源专员": {
            "category": "人力/行政",
            "core": ["招聘", "员工关系"],
            "related": ["绩效", "薪酬", "培训", "Excel", "沟通能力"],
            "bonus": ["公文写作", "行政管理", "Office"],
            "education_keywords": ["人力资源", "工商管理", "心理学", "行政管理", "劳动关系"],
        },
        "行政专员": {
            "category": "人力/行政",
            "core": ["行政管理", "公文写作"],
            "related": ["档案管理", "合同管理", "Office", "沟通能力"],
            "bonus": ["活动策划", "采购", "客户沟通"],
            "education_keywords": ["行政管理", "文秘", "工商管理", "公共管理", "中文", "汉语言"],
        },
        "市场专员": {
            "category": "市场/品牌",
            "core": ["市场调研", "品牌推广"],
            "related": ["活动策划", "客户沟通", "数据分析", "Excel", "新媒体运营"],
            "bonus": ["SEO", "SEM", "Photoshop"],
            "education_keywords": ["市场营销", "广告", "传播", "工商管理", "电子商务"],
        },
        "销售专员": {
            "category": "销售/商务",
            "core": ["销售", "客户沟通"],
            "related": ["渠道开发", "商务谈判", "市场调研", "Excel"],
            "bonus": ["品牌推广", "项目管理", "英语"],
            "education_keywords": ["市场营销", "工商管理", "国际贸易", "电子商务"],
        },
        "商务专员": {
            "category": "销售/商务",
            "core": ["商务谈判", "客户沟通"],
            "related": ["合同管理", "渠道开发", "销售", "Excel", "Office"],
            "bonus": ["英语", "项目管理", "市场调研"],
            "education_keywords": ["工商管理", "市场营销", "国际贸易", "商务英语"],
        },
        "采购专员": {
            "category": "供应链/采购",
            "core": ["采购", "供应链"],
            "related": ["合同管理", "商务谈判", "Excel", "客户沟通"],
            "bonus": ["物流", "质量管理", "英语"],
            "education_keywords": ["物流管理", "供应链", "工商管理", "国际贸易"],
        },
        "物流专员": {
            "category": "供应链/物流",
            "core": ["物流", "仓储"],
            "related": ["供应链", "Excel", "数据分析", "沟通能力"],
            "bonus": ["采购", "生产管理", "Office"],
            "education_keywords": ["物流管理", "供应链", "交通运输", "工商管理"],
        },
        "教师/助教": {
            "category": "教育/培训",
            "core": ["授课", "课程设计"],
            "related": ["教师资格证", "班级管理", "培训", "沟通能力", "Office"],
            "bonus": ["英语", "公文写作", "活动策划"],
            "education_keywords": ["师范", "教育", "汉语言", "中文", "英语", "数学", "物理", "化学"],
        },
    }

    def __init__(self, db_path=DB_PATH):
        self.db_path = db_path
        self._init_db()

    def _init_db(self):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS resumes (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                filename TEXT,
                original_filename TEXT,
                name TEXT,
                phone TEXT,
                email TEXT,
                education TEXT,
                skills TEXT,
                source_stats TEXT,
                raw_preview TEXT,
                parse_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)

        cursor.execute("PRAGMA table_info(resumes)")
        existing_columns = {row[1] for row in cursor.fetchall()}

        required_columns = {
            "original_filename": "TEXT",
            "source_stats": "TEXT",
            "raw_preview": "TEXT",
            "parse_time": "TIMESTAMP DEFAULT CURRENT_TIMESTAMP",
        }

        for col, col_type in required_columns.items():
            if col not in existing_columns:
                cursor.execute(f"ALTER TABLE resumes ADD COLUMN {col} {col_type}")
                print(f"数据库自动升级：已添加字段 {col}")

        conn.commit()
        conn.close()

    def save_to_sqlite(self, resumes):
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute("DELETE FROM resumes")

        for resume in resumes:
            cursor.execute("""
                INSERT INTO resumes
                (filename, original_filename, name, phone, email, education, skills, source_stats, raw_preview)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                resume.filename,
                resume.original_filename,
                resume.name,
                resume.phone,
                resume.email,
                json.dumps(resume.education, ensure_ascii=False),
                json.dumps(resume.skills, ensure_ascii=False),
                json.dumps(resume.source_stats, ensure_ascii=False),
                resume.raw_preview,
            ))

        conn.commit()
        conn.close()
        print(f"已保存 {len(resumes)} 条数据到 SQLite")

    def export_to_csv(self, resumes, output_path=os.path.join(EXPORT_FOLDER, "resumes.csv")):
        with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
            writer = csv.DictWriter(
                f,
                fieldnames=[
                    "original_filename",
                    "name",
                    "phone",
                    "email",
                    "education",
                    "skills",
                    "source_stats",
                    "raw_preview",
                ],
            )
            writer.writeheader()

            for resume in resumes:
                row = resume.to_dict()
                writer.writerow({
                    "original_filename": row["original_filename"],
                    "name": row["name"],
                    "phone": row["phone"],
                    "email": row["email"],
                    "education": json.dumps(row["education"], ensure_ascii=False),
                    "skills": json.dumps(row["skills"], ensure_ascii=False),
                    "source_stats": json.dumps(row["source_stats"], ensure_ascii=False),
                    "raw_preview": row["raw_preview"],
                })

        print(f"已导出 CSV: {output_path}")

    def export_to_json(self, resumes, output_path=os.path.join(EXPORT_FOLDER, "resumes.json")):
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump([r.to_dict() for r in resumes], f, ensure_ascii=False, indent=4)

        print(f"已导出 JSON: {output_path}")

    def skill_statistics(self, resumes):
        counter = Counter()
        for resume in resumes:
            counter.update(resume.skills)
        return counter

    def overview_statistics(self, resumes, matching_result):
        total = len(resumes)
        skill_counter = self.skill_statistics(resumes)
        avg_skills = round(sum(len(r.skills) for r in resumes) / total, 1) if total else 0
        table_files = sum(1 for r in resumes if r.source_stats.get("tables", 0) > 0)
        failed_files = sum(1 for r in resumes if r.source_stats.get("chars", 0) == 0)
        named_people = sum(1 for r in resumes if r.name)
        displayed_people = sum(1 for p in matching_result if p["top_jobs"])

        return {
            "total": total,
            "named_people": named_people,
            "unique_skills": len(skill_counter),
            "avg_skills": avg_skills,
            "table_files": table_files,
            "failed_files": failed_files,
            "displayed_people": displayed_people,
            "top_skill": skill_counter.most_common(1)[0][0] if skill_counter else "无",
            "job_count": len(self.JOB_PROFILES),
        }

    def generate_visual_reports(self, resumes):
        charts = {}

        if self.plot_skill_dist(resumes, os.path.join(STATIC_FOLDER, "skill_dist.png")):
            charts["skill"] = "skill_dist.png"

        if self.plot_job_match_heatmap(resumes, os.path.join(STATIC_FOLDER, "job_match_heatmap.png")):
            charts["heatmap"] = "job_match_heatmap.png"

        return charts

    def plot_skill_dist(self, resumes, save_path):
        counter = self.skill_statistics(resumes)

        if not counter:
            return False

        items = counter.most_common(25)
        labels = [x[0] for x in items][::-1]
        values = [x[1] for x in items][::-1]

        height = max(8, len(labels) * 0.55 + 2)
        plt.figure(figsize=(14, height))
        bars = plt.barh(labels, values, color="#2563EB", alpha=0.88, height=0.58)

        plt.title("Top 技能/能力出现频次", fontsize=24, pad=22, weight="bold")
        plt.xlabel("出现次数", fontsize=16, labelpad=12)
        plt.ylabel("技能/能力名称", fontsize=16, labelpad=12)
        plt.xticks(fontsize=13)
        plt.yticks(fontsize=14)
        plt.grid(axis="x", linestyle="--", alpha=0.25)

        ax = plt.gca()
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
        ax.spines["left"].set_alpha(0.25)
        ax.spines["bottom"].set_alpha(0.25)

        max_value = max(values) if values else 1
        plt.xlim(0, max_value + max(1, max_value * 0.18))

        for bar in bars:
            width = bar.get_width()
            plt.text(
                width + max(0.05, max_value * 0.015),
                bar.get_y() + bar.get_height() / 2,
                str(int(width)),
                va="center",
                fontsize=13,
                weight="bold",
                color="#111827",
            )

        plt.subplots_adjust(left=0.22, right=0.95, top=0.88, bottom=0.12)
        plt.savefig(save_path, dpi=180, bbox_inches="tight")
        plt.close()

        return True

    def plot_job_match_heatmap(self, resumes, save_path):
        valid_resumes = [r for r in resumes if r.skills or r.education]

        if not valid_resumes:
            return False

        display_resumes = valid_resumes[:30]
        names = [(r.name or os.path.splitext(r.original_filename)[0])[:14] for r in display_resumes]
        jobs = list(self.JOB_PROFILES.keys())
        matrix = []

        for resume in display_resumes:
            row = []
            for job_name in jobs:
                score_info = self._score_job_match(resume, self.JOB_PROFILES[job_name])
                row.append(score_info["score"])
            matrix.append(row)

        data = np.array(matrix)

        fig_width = max(22, len(jobs) * 0.95)
        fig_height = max(8, len(names) * 0.55 + 3)

        plt.figure(figsize=(fig_width, fig_height))

        im = plt.imshow(data, cmap="YlGnBu", aspect="auto", vmin=0, vmax=100)
        cbar = plt.colorbar(im, fraction=0.018, pad=0.012)
        cbar.set_label("岗位匹配分", fontsize=15, labelpad=12)
        cbar.ax.tick_params(labelsize=12)

        plt.xticks(range(len(jobs)), jobs, rotation=38, ha="right", fontsize=11)
        plt.yticks(range(len(names)), names, fontsize=12)
        plt.title("候选人与多类型岗位匹配热力图", fontsize=24, pad=24, weight="bold")

        ax = plt.gca()
        ax.set_xlabel("推荐岗位", fontsize=14, labelpad=16)
        ax.set_ylabel("候选人", fontsize=14, labelpad=16)
        ax.tick_params(axis="both", length=0)

        plt.subplots_adjust(left=0.10, right=0.96, top=0.88, bottom=0.34)
        plt.savefig(save_path, dpi=200, bbox_inches="tight")
        plt.close()

        return True

    def job_matching(self, resumes):
        results = []

        for resume in resumes:
            matches = []

            for job_name, profile in self.JOB_PROFILES.items():
                score_info = self._score_job_match(resume, profile)

                # 关键修改：推荐列表不显示 0 分，也不显示接近 0 分的无意义岗位。
                if score_info["score"] < self.MIN_DISPLAY_SCORE:
                    continue

                matches.append({
                    "job_name": job_name,
                    "category": profile.get("category", "综合"),
                    "matched_core_skills": score_info["matched_core"],
                    "matched_related_skills": score_info["matched_related"],
                    "matched_bonus_skills": score_info["matched_bonus"],
                    "matched_education_keywords": score_info["matched_education_keywords"],
                    "missing_core_skills": score_info["missing_core"][:5],
                    "match_rate": round(score_info["score"], 1),
                    "core_rate": round(score_info["core_rate"], 1),
                    "related_rate": round(score_info["related_rate"], 1),
                    "recommend_level": "待定",
                    "reason": "",
                })

            matches = sorted(
                matches,
                key=lambda x: (
                    x["match_rate"],
                    len(x["matched_core_skills"]),
                    len(x["matched_related_skills"]),
                    len(x["matched_education_keywords"]),
                ),
                reverse=True,
            )[:3]

            for rank, item in enumerate(matches, start=1):
                item["recommend_level"] = self._relative_recommend_level(
                    resume=resume,
                    rank=rank,
                    score=item["match_rate"],
                    core_hit=len(item["matched_core_skills"]),
                    related_hit=len(item["matched_related_skills"]),
                    edu_hit=len(item["matched_education_keywords"]),
                )
                item["reason"] = self._recommend_reason(item)

            # 关键修改：该候选人如果没有任何有效匹配岗位，就不进入推荐表。
            if matches:
                results.append({
                    "name": resume.name or "未识别姓名",
                    "filename": resume.original_filename,
                    "top_jobs": matches,
                })

        return results

    def _score_job_match(self, resume, profile):
        candidate_skills = resume.skills or []
        candidate_norms = {normalize_token(s) for s in candidate_skills}
        edu_text = " ".join(resume.education or [])

        core = profile.get("core", [])
        related = profile.get("related", [])
        bonus = profile.get("bonus", [])
        education_keywords = profile.get("education_keywords", [])

        matched_core = [s for s in core if normalize_token(s) in candidate_norms]
        matched_related = [s for s in related if normalize_token(s) in candidate_norms]
        matched_bonus = [s for s in bonus if normalize_token(s) in candidate_norms]
        matched_education_keywords = [k for k in education_keywords if k and k in edu_text]

        missing_core = [s for s in core if normalize_token(s) not in candidate_norms]

        core_rate = len(matched_core) / len(core) * 100 if core else 0
        related_rate = len(matched_related) / len(related) * 100 if related else 0
        bonus_rate = min(len(matched_bonus) / max(1, len(bonus)), 1) * 100 if bonus else 0
        edu_rate = min(len(matched_education_keywords) / max(1, len(education_keywords)), 1) * 100 if education_keywords else 0
        skill_breadth = min(len(candidate_skills) / 8, 1) * 100

        score = (
            core_rate * 0.38
            + related_rate * 0.30
            + bonus_rate * 0.12
            + edu_rate * 0.12
            + skill_breadth * 0.08
        )

        if len(matched_core) > 0:
            score += 8
        elif len(matched_related) >= 2:
            score += 6
        elif len(matched_related) == 1:
            score += 3

        if matched_education_keywords:
            score += 6

        if not candidate_skills and matched_education_keywords:
            score = max(score, 32)

        if not candidate_skills and not resume.education:
            score = 0

        return {
            "score": max(0, min(score, 100)),
            "core_rate": core_rate,
            "related_rate": related_rate,
            "bonus_rate": bonus_rate,
            "edu_rate": edu_rate,
            "matched_core": matched_core,
            "matched_related": matched_related,
            "matched_bonus": matched_bonus,
            "matched_education_keywords": matched_education_keywords,
            "missing_core": missing_core,
        }

    def _relative_recommend_level(self, resume, rank, score, core_hit, related_hit, edu_hit):
        has_info = bool(resume.skills or resume.education)

        if not has_info:
            return "信息不足"

        if rank == 1:
            if score >= 60 or core_hit >= 2:
                return "优先推荐"
            if score >= 25 or core_hit >= 1 or related_hit >= 1 or edu_hit >= 1:
                return "推荐"
            return "可培养"

        if rank == 2:
            if score >= 50 or core_hit >= 2:
                return "推荐"
            return "可培养"

        return "可培养"

    def _recommend_reason(self, item):
        parts = []

        if item["matched_core_skills"]:
            parts.append("核心能力命中：" + "、".join(item["matched_core_skills"]))

        if item["matched_related_skills"]:
            parts.append("相关能力命中：" + "、".join(item["matched_related_skills"][:5]))

        if item["matched_education_keywords"]:
            parts.append("专业背景相关：" + "、".join(item["matched_education_keywords"]))

        if item["matched_bonus_skills"]:
            parts.append("加分能力：" + "、".join(item["matched_bonus_skills"][:4]))

        if not parts:
            parts.append("当前简历信息较少，系统根据已有文本给出相对最接近岗位")

        return "；".join(parts)


# ===================== Flask Web 界面 =====================
app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["STATIC_FOLDER"] = STATIC_FOLDER
app.config["EXPORT_FOLDER"] = EXPORT_FOLDER


INDEX_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>智能简历解析器</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; font-family: "Microsoft YaHei", sans-serif; }
        body { min-height: 100vh; background: linear-gradient(135deg, #eef2ff, #f8fafc, #ecfdf5); padding: 48px 18px; color: #1f2937; }
        .container { max-width: 920px; margin: 0 auto; background: rgba(255,255,255,0.94); padding: 42px; border-radius: 24px; box-shadow: 0 24px 80px rgba(15, 23, 42, 0.12); }
        .badge { display: inline-block; background: #ecfdf5; color: #047857; padding: 8px 14px; border-radius: 999px; font-size: 13px; font-weight: 700; margin-bottom: 18px; }
        h1 { font-size: 34px; color: #111827; margin-bottom: 12px; }
        .subtitle { color: #6b7280; line-height: 1.8; margin-bottom: 28px; }
        .upload-area { width: 100%; height: 230px; border: 2.5px dashed #b7c2d0; border-radius: 22px; display: flex; flex-direction: column; justify-content: center; align-items: center; cursor: pointer; transition: all 0.25s ease; margin-bottom: 20px; background: #f8fafc; }
        .upload-area:hover, .upload-area.dragover { border-color: #2563eb; background: #eff6ff; transform: translateY(-2px); }
        .upload-icon { font-size: 48px; margin-bottom: 14px; }
        .upload-text { font-size: 20px; color: #111827; font-weight: 700; margin-bottom: 8px; }
        .upload-tip { font-size: 14px; color: #6b7280; }
        #fileInput { display: none; }
        .file-list { width: 100%; min-height: 72px; border: 1px solid #e5e7eb; border-radius: 18px; padding: 14px; margin-bottom: 24px; background: #ffffff; }
        .file-item { display: flex; justify-content: space-between; align-items: center; padding: 12px 14px; background: #f9fafb; border-radius: 12px; margin-bottom: 10px; border: 1px solid #eef2f7; }
        .file-name { color: #374151; font-size: 14px; word-break: break-all; }
        .remove-btn { color: #ef4444; cursor: pointer; font-size: 20px; font-weight: bold; padding: 2px 8px; }
        .btn { display: block; width: 230px; height: 52px; border: none; border-radius: 14px; font-size: 16px; font-weight: 700; cursor: pointer; margin: 0 auto; background: linear-gradient(135deg, #2563eb, #10b981); color: white; box-shadow: 0 10px 24px rgba(37, 99, 235, 0.22); }
        .btn:disabled { background: #cbd5e1; cursor: not-allowed; box-shadow: none; }
        .tips { margin-top: 24px; background: #fffbeb; border: 1px solid #fde68a; color: #92400e; padding: 14px 16px; border-radius: 14px; font-size: 14px; line-height: 1.8; }
    </style>
</head>
<body>
    <div class="container">
        <div class="badge">多专业岗位推荐版 v3</div>
        <h1>智能简历解析器</h1>
        <div class="subtitle">批量上传 PDF、DOCX、TXT 简历，自动提取姓名、电话、邮箱、教育经历、技能，并根据多专业岗位库给出有效推荐。</div>

        <div class="upload-area" id="uploadArea">
            <div class="upload-icon">📄</div>
            <div class="upload-text">拖拽简历文件到此处，或点击选择</div>
            <div class="upload-tip">支持 PDF、DOCX、TXT，可批量上传</div>
        </div>

        <div class="file-list" id="fileList">
            <div style="text-align: center; color: #9ca3af; font-size: 14px; padding: 12px;">暂无上传文件</div>
        </div>

        <form id="uploadForm" enctype="multipart/form-data" action="/parse" method="post">
            <input type="file" id="fileInput" name="files[]" multiple accept=".pdf,.txt,.docx">
            <button type="submit" class="btn" id="submitBtn" disabled>开始智能解析</button>
        </form>

    </div>

    <script>
        const uploadArea = document.getElementById('uploadArea');
        const fileInput = document.getElementById('fileInput');
        const fileList = document.getElementById('fileList');
        const submitBtn = document.getElementById('submitBtn');
        let selectedFiles = [];

        function updateFileInput() {
            const dt = new DataTransfer();
            selectedFiles.forEach(file => dt.items.add(file));
            fileInput.files = dt.files;
        }

        function updateFileList() {
            fileList.innerHTML = '';

            if (selectedFiles.length === 0) {
                fileList.innerHTML = '<div style="text-align: center; color: #9ca3af; font-size: 14px; padding: 12px;">暂无上传文件</div>';
                submitBtn.disabled = true;
                updateFileInput();
                return;
            }

            submitBtn.disabled = false;

            selectedFiles.forEach((file, index) => {
                const item = document.createElement('div');
                item.className = 'file-item';
                item.innerHTML = `<span class="file-name">${file.name}</span><span class="remove-btn" data-index="${index}">×</span>`;
                fileList.appendChild(item);
            });

            document.querySelectorAll('.remove-btn').forEach(btn => {
                btn.addEventListener('click', e => {
                    selectedFiles.splice(parseInt(e.target.dataset.index), 1);
                    updateFileList();
                });
            });

            updateFileInput();
        }

        function addFiles(files) {
            const allowExt = ['pdf', 'txt', 'docx'];
            for (let i = 0; i < files.length; i++) {
                const ext = files[i].name.split('.').pop().toLowerCase();
                if (allowExt.includes(ext)) selectedFiles.push(files[i]);
            }
            updateFileList();
        }

        uploadArea.addEventListener('click', () => fileInput.click());
        fileInput.addEventListener('change', e => addFiles(e.target.files));
        uploadArea.addEventListener('dragover', e => {
            e.preventDefault();
            uploadArea.classList.add('dragover');
        });
        uploadArea.addEventListener('dragleave', () => uploadArea.classList.remove('dragover'));
        uploadArea.addEventListener('drop', e => {
            e.preventDefault();
            uploadArea.classList.remove('dragover');
            addFiles(e.dataTransfer.files);
        });
    </script>
</body>
</html>
"""


RESULT_TEMPLATE = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>解析结果</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; font-family: "Microsoft YaHei", sans-serif; }
        body { background: #f3f6fb; color: #1f2937; padding: 28px 16px 48px; }
        .container { max-width: 1680px; margin: 0 auto; }
        .hero { background: linear-gradient(135deg, #2563eb, #10b981); color: white; border-radius: 26px; padding: 30px 34px; margin-bottom: 24px; box-shadow: 0 20px 60px rgba(37, 99, 235, 0.22); }
        .hero-top { display: flex; justify-content: space-between; align-items: flex-start; gap: 18px; flex-wrap: wrap; }
        .badge { display: inline-block; padding: 7px 13px; border-radius: 999px; background: rgba(255,255,255,0.18); border: 1px solid rgba(255,255,255,0.32); font-size: 13px; font-weight: 700; margin-bottom: 14px; }
        h1 { font-size: 34px; margin-bottom: 10px; letter-spacing: 0.5px; }
        .hero-desc { font-size: 15px; line-height: 1.8; color: rgba(255,255,255,0.88); }
        .hero-actions { display: flex; gap: 10px; flex-wrap: wrap; }
        .btn { display: inline-block; padding: 11px 16px; border-radius: 13px; text-decoration: none; font-size: 14px; font-weight: 800; transition: all 0.2s ease; }
        .btn-light { background: white; color: #1d4ed8; }
        .btn-ghost { background: rgba(255,255,255,0.16); color: white; border: 1px solid rgba(255,255,255,0.3); }
        .btn:hover { transform: translateY(-1px); }
        .stats-grid { display: grid; grid-template-columns: repeat(auto-fit, minmax(150px, 1fr)); gap: 16px; margin-bottom: 24px; }
        .stat-card { background: white; padding: 19px 18px; border-radius: 20px; box-shadow: 0 10px 30px rgba(15,23,42,0.06); border: 1px solid #edf2f7; }
        .stat-value { font-size: 25px; font-weight: 900; color: #111827; margin-bottom: 7px; word-break: break-all; }
        .stat-label { color: #6b7280; font-size: 13px; }
        .section { background: white; border-radius: 24px; padding: 26px; margin-bottom: 24px; box-shadow: 0 10px 30px rgba(15,23,42,0.06); border: 1px solid #edf2f7; }
        .section-title-row { display: flex; justify-content: space-between; align-items: center; gap: 12px; flex-wrap: wrap; margin-bottom: 20px; }
        .section-title { font-size: 21px; color: #111827; border-left: 5px solid #2563eb; padding-left: 12px; font-weight: 900; }
        .section-tip { color: #6b7280; font-size: 13px; }
        .resume-list { display: grid; grid-template-columns: repeat(auto-fill, minmax(360px, 1fr)); gap: 18px; }
        .resume-card { border: 1px solid #e5e7eb; border-radius: 20px; padding: 18px; background: linear-gradient(180deg, #ffffff, #fbfdff); transition: all 0.2s ease; }
        .resume-card:hover { transform: translateY(-2px); box-shadow: 0 14px 36px rgba(15,23,42,0.08); }
        .resume-head { display: flex; justify-content: space-between; align-items: flex-start; gap: 12px; margin-bottom: 12px; }
        .resume-name { font-size: 20px; font-weight: 900; color: #111827; margin-bottom: 6px; }
        .filename { font-size: 12px; color: #94a3b8; word-break: break-all; }
        .status-badge { flex-shrink: 0; padding: 5px 9px; border-radius: 999px; font-size: 12px; font-weight: 800; }
        .status-ok { background: #ecfdf5; color: #047857; }
        .status-warn { background: #fff7ed; color: #c2410c; }
        .info-grid { display: grid; grid-template-columns: 74px minmax(0, 1fr); gap: 8px 10px; font-size: 14px; line-height: 1.7; margin-top: 12px; }
        .info-label { color: #64748b; font-weight: 700; }
        .info-value { color: #374151; word-break: break-word; }
        .tags { display: flex; flex-wrap: wrap; gap: 7px; margin-top: 2px; }
        .tag { background: #eff6ff; color: #1d4ed8; padding: 4px 9px; border-radius: 999px; font-size: 12px; font-weight: 700; }
        .tag-core { background: #ecfdf5; color: #047857; }
        .tag-bonus { background: #fef3c7; color: #92400e; }
        .tag-edu { background: #f5f3ff; color: #6d28d9; }
        .mini { color: #64748b; font-size: 12px; margin-top: 12px; line-height: 1.6; background: #f8fafc; padding: 9px 10px; border-radius: 12px; }
        .warning { margin-top: 10px; padding: 9px 11px; background: #fff7ed; color: #9a3412; border: 1px solid #fed7aa; border-radius: 12px; font-size: 12px; line-height: 1.6; }
        .preview { margin-top: 10px; padding: 10px 11px; background: #f8fafc; border-radius: 12px; font-size: 12px; color: #64748b; max-height: 120px; overflow: auto; white-space: pre-wrap; border: 1px solid #edf2f7; }
        .match-table-wrapper { overflow-x: auto; border-radius: 18px; border: 1px solid #edf2f7; }
        .match-table { border-collapse: collapse; width: 100%; min-width: 1320px; background: white; }
        .match-table th, .match-table td { border-bottom: 1px solid #edf2f7; padding: 14px; text-align: left; vertical-align: top; font-size: 14px; }
        .match-table th { background: #f8fafc; font-weight: 900; color: #374151; }
        .match-table tr:last-child td { border-bottom: none; }
        .level-strong { color: #059669; font-weight: 900; }
        .level-normal { color: #2563eb; font-weight: 900; }
        .level-train { color: #d97706; font-weight: 900; }
        .level-info { color: #64748b; font-weight: 900; }
        .progress { width: 150px; height: 9px; background: #e5e7eb; border-radius: 999px; overflow: hidden; margin-top: 7px; }
        .progress-inner { height: 100%; background: linear-gradient(90deg, #2563eb, #10b981); border-radius: 999px; }

        /* 关键修改：两个可视化图改为上下两行，每个图独占一行 */
        .charts-layout { display: grid; grid-template-columns: 1fr; gap: 24px; align-items: stretch; }
        .chart-card { border: 1px solid #edf2f7; border-radius: 20px; padding: 18px; background: #fbfdff; overflow: auto; }
        .chart-title { font-size: 17px; font-weight: 900; margin-bottom: 14px; color: #111827; }
        .chart-card img { width: 100%; max-width: none; min-width: 820px; border-radius: 15px; display: block; background: white; }

        .empty { color: #94a3b8; text-align: center; padding: 22px; font-size: 14px; }
        @media (max-width: 1000px) {
            .hero { padding: 26px 22px; }
            h1 { font-size: 28px; }
            .resume-list { grid-template-columns: 1fr; }
            .chart-card img { min-width: 680px; }
        }
    </style>
</head>

<body>
    <div class="container">

        <div class="hero">
            <div class="hero-top">
                <div>
                    <div class="badge">解析完成</div>
                    <h1>简历解析结果</h1>
                </div>

                <div class="hero-actions">
                    <a href="/" class="btn btn-ghost">← 返回上传</a>
                    <a href="/download/resumes.csv" class="btn btn-light">下载 CSV</a>
                    <a href="/download/resumes.json" class="btn btn-light">下载 JSON</a>
                </div>
            </div>
        </div>

        <div class="stats-grid">
            <div class="stat-card"><div class="stat-value">{{ overview.total }}</div><div class="stat-label">处理文件数</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.named_people }}</div><div class="stat-label">识别到姓名</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.unique_skills }}</div><div class="stat-label">技能/能力种类</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.avg_skills }}</div><div class="stat-label">人均能力数</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.job_count }}</div><div class="stat-label">岗位库数量</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.displayed_people }}</div><div class="stat-label">有有效推荐人数</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.failed_files }}</div><div class="stat-label">未提取文本文件</div></div>
            <div class="stat-card"><div class="stat-value">{{ overview.top_skill }}</div><div class="stat-label">最高频能力</div></div>
        </div>

        <div class="section">
            <div class="section-title-row">
                <div class="section-title">解析结果明细</div>
                <div class="section-tip">展示每份简历解析出的基础信息、教育经历和技能能力</div>
            </div>

            {% if resumes|length == 0 %}
            <div class="empty">没有文件被处理。</div>
            {% endif %}

            <div class="resume-list">
                {% for resume in resumes %}
                <div class="resume-card">
                    <div class="resume-head">
                        <div>
                            <div class="resume-name">{{ resume.name or '未识别姓名' }}</div>
                            <div class="filename">{{ resume.original_filename }}</div>
                        </div>

                        {% if resume.source_stats.chars > 0 and resume.name %}
                        <div class="status-badge status-ok">已识别</div>
                        {% else %}
                        <div class="status-badge status-warn">待核对</div>
                        {% endif %}
                    </div>

                    <div class="info-grid">
                        <div class="info-label">电话</div>
                        <div class="info-value">{{ resume.phone or '未提取到' }}</div>

                        <div class="info-label">邮箱</div>
                        <div class="info-value">{{ resume.email or '未提取到' }}</div>

                        <div class="info-label">教育</div>
                        <div class="info-value">{{ resume.education|join('、') or '未提取到' }}</div>

                        <div class="info-label">能力</div>
                        <div class="info-value">
                            <div class="tags">
                                {% for skill in resume.skills %}
                                <span class="tag">{{ skill }}</span>
                                {% endfor %}
                                {% if resume.skills|length == 0 %}
                                <span class="mini">未提取到</span>
                                {% endif %}
                            </div>
                        </div>
                    </div>

                    <div class="mini">
                        解析器：{{ resume.source_stats.parser or '未知' }}；
                        字符数：{{ resume.source_stats.chars }}；
                        表格数：{{ resume.source_stats.tables }}
                        {% if resume.source_stats.ocr_used %}；已使用 OCR{% endif %}
                    </div>

                    {% if resume.source_stats.warning %}
                    <div class="warning">{{ resume.source_stats.warning }}</div>
                    {% endif %}

                    {% if resume.raw_preview %}
                    <div class="preview">文本预览：{{ resume.raw_preview }}</div>
                    {% endif %}
                </div>
                {% endfor %}
            </div>
        </div>

        <div class="section">
            <div class="section-title-row">
                <div class="section-title">多专业岗位推荐 Top 3</div>
                <div class="section-tip">已过滤 0 分岗位；若某候选人完全没有有效命中，则不在此表展示</div>
            </div>

            {% if matching_result|length == 0 %}
            <div class="empty">暂无有效岗位推荐。请检查简历是否成功提取到教育经历、专业信息或技能关键词。</div>
            {% else %}
            <div class="match-table-wrapper">
                <table class="match-table">
                    <thead>
                        <tr>
                            <th>姓名</th>
                            <th>推荐岗位</th>
                            <th>岗位类别</th>
                            <th>核心命中</th>
                            <th>相关命中</th>
                            <th>加分能力</th>
                            <th>专业背景</th>
                            <th>匹配分</th>
                            <th>推荐等级</th>
                            <th>推荐理由</th>
                        </tr>
                    </thead>

                    <tbody>
                        {% for person in matching_result %}
                            {% for job in person.top_jobs %}
                            <tr>
                                {% if loop.first %}
                                <td rowspan="{{ person.top_jobs|length }}">{{ person.name }}</td>
                                {% endif %}

                                <td>{{ job.job_name }}</td>
                                <td>{{ job.category }}</td>

                                <td>
                                    <div class="tags">
                                        {% for skill in job.matched_core_skills %}
                                        <span class="tag tag-core">{{ skill }}</span>
                                        {% endfor %}
                                        {% if job.matched_core_skills|length == 0 %}暂无{% endif %}
                                    </div>
                                </td>

                                <td>
                                    <div class="tags">
                                        {% for skill in job.matched_related_skills %}
                                        <span class="tag">{{ skill }}</span>
                                        {% endfor %}
                                        {% if job.matched_related_skills|length == 0 %}暂无{% endif %}
                                    </div>
                                </td>

                                <td>
                                    <div class="tags">
                                        {% for skill in job.matched_bonus_skills %}
                                        <span class="tag tag-bonus">{{ skill }}</span>
                                        {% endfor %}
                                        {% if job.matched_bonus_skills|length == 0 %}暂无{% endif %}
                                    </div>
                                </td>

                                <td>
                                    <div class="tags">
                                        {% for kw in job.matched_education_keywords %}
                                        <span class="tag tag-edu">{{ kw }}</span>
                                        {% endfor %}
                                        {% if job.matched_education_keywords|length == 0 %}暂无{% endif %}
                                    </div>
                                </td>

                                <td>
                                    {{ job.match_rate }}分
                                    <div class="progress">
                                        <div class="progress-inner" style="width: {{ job.match_rate }}%;"></div>
                                    </div>
                                </td>

                                {% if job.recommend_level == '优先推荐' %}
                                <td class="level-strong">{{ job.recommend_level }}</td>
                                {% elif job.recommend_level == '推荐' %}
                                <td class="level-normal">{{ job.recommend_level }}</td>
                                {% elif job.recommend_level == '可培养' %}
                                <td class="level-train">{{ job.recommend_level }}</td>
                                {% else %}
                                <td class="level-info">{{ job.recommend_level }}</td>
                                {% endif %}

                                <td>{{ job.reason }}</td>
                            </tr>
                            {% endfor %}
                        {% endfor %}
                    </tbody>
                </table>
            </div>
            {% endif %}
        </div>

        {% if charts %}
        <div class="section">
            <div class="section-title-row">
                <div class="section-title">可视化分析报告</div>
            </div>

            <div class="charts-layout">
                {% if charts.skill %}
                <div class="chart-card">
                    <div class="chart-title">技能/能力出现频次</div>
                    <img src="/static/{{ charts.skill }}?v={{ ts }}" alt="技能分布">
                </div>
                {% endif %}

                {% if charts.heatmap %}
                <div class="chart-card">
                    <div class="chart-title">多专业岗位匹配热力图</div>
                    <img src="/static/{{ charts.heatmap }}?v={{ ts }}" alt="岗位匹配热力图">
                </div>
                {% endif %}
            </div>
        </div>
        {% endif %}

    </div>
</body>
</html>
"""


@app.route("/")
def index():
    return render_template_string(INDEX_TEMPLATE)


@app.route("/parse", methods=["POST"])
def parse_resume():
    for f in os.listdir(UPLOAD_FOLDER):
        path = os.path.join(UPLOAD_FOLDER, f)
        if os.path.isfile(path):
            os.remove(path)

    files = request.files.getlist("files[]")

    if not files or files[0].filename == "":
        return redirect(url_for("index"))

    filename_map = {}
    saved_count = 0

    for file in files:
        if not file or not file.filename:
            continue

        original_name = safe_original_filename(file.filename)

        if not allowed_file(original_name):
            print(f"跳过不支持文件: {original_name}")
            continue

        ext = original_name.rsplit(".", 1)[-1].lower()
        safe_base = secure_filename(os.path.splitext(original_name)[0]) or "resume"
        save_name = f"{safe_base}_{uuid.uuid4().hex[:8]}.{ext}"

        file.save(os.path.join(app.config["UPLOAD_FOLDER"], save_name))
        filename_map[save_name] = original_name
        saved_count += 1

    if saved_count == 0:
        return redirect(url_for("index"))

    parser = Parser(UPLOAD_FOLDER, filename_map=filename_map)
    resumes = parser.batch_parse()

    report_gen = ReportGenerator()
    report_gen.save_to_sqlite(resumes)
    report_gen.export_to_csv(resumes)
    report_gen.export_to_json(resumes)

    matching_result = report_gen.job_matching(resumes)
    charts = report_gen.generate_visual_reports(resumes)
    overview = report_gen.overview_statistics(resumes, matching_result)

    return render_template_string(
        RESULT_TEMPLATE,
        resumes=[r.to_dict() for r in resumes],
        matching_result=matching_result,
        charts=charts,
        overview=overview,
        ts=int(datetime.now().timestamp()),
    )


@app.route("/download/<path:filename>")
def download_file(filename):
    if filename not in {"resumes.csv", "resumes.json"}:
        return "不允许下载该文件", 403

    return send_from_directory(EXPORT_FOLDER, filename, as_attachment=True)


def open_browser():
    webbrowser.open("http://127.0.0.1:5000")


if __name__ == "__main__":
    print("=" * 70)
    print("智能简历解析器启动中...")
    print("访问地址：http://127.0.0.1:5000")
    print("基础依赖：pip install flask python-docx PyPDF2 matplotlib numpy werkzeug")
    print("推荐依赖：pip install pdfplumber")
    print("扫描版PDF额外依赖：pip install pytesseract pdf2image pillow")
    print("扫描版PDF还需安装 Tesseract OCR 和 Poppler")
    print("=" * 70)

    threading.Timer(1.0, open_browser).start()
    app.run(debug=True, use_reloader=False)
